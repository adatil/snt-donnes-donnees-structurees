#!/usr/bin/env python3
import pandas as pd
from docx import Document
import json

# Lire le fichier Excel
print("=== LECTURE DU FICHIER EXCEL ===")
try:
    df = pd.read_excel("Liste des films ayant réalisé plus d'un million d'entrées -.xlsx")
    print(f"Nombre de lignes: {len(df)}")
    print(f"Colonnes: {list(df.columns)}")
    print("\nPremières lignes:")
    print(df.head(20))
    print("\nDernières lignes:")
    print(df.tail(20))

    # Sauvegarder en CSV
    df.to_csv("films.csv", index=False)
    print("\nFichier sauvegardé en films.csv")

    # Statistiques sur les années
    if 'Année' in df.columns or 'année' in df.columns or 'Annee' in df.columns:
        year_col = [c for c in df.columns if 'ann' in c.lower()][0] if any('ann' in c.lower() for c in df.columns) else None
        if year_col:
            print(f"\nStatistiques sur les années (colonne: {year_col}):")
            print(df[year_col].value_counts().sort_index(ascending=False).head(20))
except Exception as e:
    print(f"Erreur lors de la lecture du fichier Excel: {e}")

# Lire le fichier TP
print("\n\n=== LECTURE DU TP ===")
try:
    doc = Document("TP sur les données structurées-1.docx")
    print("Contenu du TP:")
    print("=" * 80)
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)

    # Sauvegarder en texte
    with open("tp_original.txt", "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")
    print("\n\nFichier TP sauvegardé en tp_original.txt")
except Exception as e:
    print(f"Erreur lors de la lecture du TP: {e}")

# Lire le fichier correction
print("\n\n=== LECTURE DE LA CORRECTION ===")
try:
    doc = Document("Tableur et tris-correctiontp1.docx")
    print("Contenu de la correction:")
    print("=" * 80)
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)

    # Sauvegarder en texte
    with open("correction_original.txt", "w", encoding="utf-8") as f:
        for para in doc.paragraphs:
            f.write(para.text + "\n")
    print("\n\nFichier correction sauvegardé en correction_original.txt")
except Exception as e:
    print(f"Erreur lors de la lecture de la correction: {e}")
